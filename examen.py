import mysql.connector
import argparse
import re
from docx import Document
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import os

class SistemaVentas:
    def __init__(self):
      
        self.conexion = mysql.connector.connect(
            host="localhost",
            user="root",
            password="",
            database="algoritmos"
        )
        self.cursor = self.conexion.cursor()

    def cerrar_conexion(self):
        self.cursor.close()
        self.conexion.close()

    # Control de Inventario
    def listar_vehiculos(self):
        self.cursor.execute("SELECT codigo, nombre, existencia, proveedor, precio FROM inventario")
        for (codigo, nombre, existencia, proveedor, precio) in self.cursor:
            print(f"Código: {codigo}, Nombre: {nombre}, Existencia: {existencia}, Proveedor: {proveedor}, Precio: {precio}")

    def crear_vehiculo(self, codigo, nombre, existencia, proveedor, precio):
        query = "INSERT INTO inventario (codigo, nombre, existencia, proveedor, precio) VALUES (%s, %s, %s, %s, %s)"
        values = (codigo, nombre, existencia, proveedor, precio)
        self.cursor.execute(query, values)
        self.conexion.commit()
        print(f"vehiculo '{nombre}' creado con éxito.")

    def actualizar_vehiculo(self, codigo, nombre, existencia, proveedor, precio):
        query = "UPDATE inventario SET nombre = %s, existencia = %s, proveedor = %s, precio = %s WHERE codigo = %s"
        values = (nombre, existencia, proveedor, precio, codigo)
        self.cursor.execute(query, values)
        self.conexion.commit()
        print(f"vehiculo '{nombre}' actualizado con éxito.")

    def editar_existencias(self, codigo, cantidad):
        query = "UPDATE inventario SET existencia = existencia + %s WHERE codigo = %s"
        values = (cantidad, codigo)
        self.cursor.execute(query, values)
        self.conexion.commit()
        print(f"Existencias del vehiculo editadas con éxito.")

    def eliminar_vehiculo(self, codigo):
        query = "DELETE FROM inventario WHERE codigo = %s"
        self.cursor.execute(query, (codigo,))
        self.conexion.commit()
        print(f"vehiculo eliminado con éxito.")

    # Control de vehiculos
    def listar_vehiculo(self):
        self.cursor.execute("SELECT codigo, nombre, direccion FROM clientes")
        for (codigo, nombre, direccion) in self.cursor:
            print(f"Código: {codigo}, Nombre: {nombre}, Dirección: {direccion}")

    def crear_vehiculo(self, codigo, nombre, direccion):
        query = "INSERT INTO clientes (codigo, nombre, direccion) VALUES (%s, %s, %s)"
        values = (codigo, nombre, direccion)
        self.cursor.execute(query, values)
        self.conexion.commit()
        print(f"Cliente '{nombre}' creado con éxito.")

    def editar_vehiculo(self, codigo, nombre, direccion):
        query = "UPDATE clientes SET nombre = %s, direccion = %s WHERE codigo = %s"
        values = (nombre, direccion, codigo)
        self.cursor.execute(query, values)
        self.conexion.commit()
        print(f"Cliente '{nombre}' actualizado con éxito.")

    def eliminar_vehiculo(self, codigo):
        query = "DELETE FROM clientes WHERE codigo = %s"
        self.cursor.execute(query, (codigo,))
        self.conexion.commit()
        print(f"Cliente eliminado con éxito.")

    # Control de Ventas
    def listar_vehiculo(self):
        self.cursor.execute("SELECT codigo_vehiculo, codigo_cliente, cantidad, total FROM ventas")
        for (codigo_vehiculo, codigo_cliente, cantidad, total) in self.cursor:
            print(f"Código de vehiculo: {codigo_vehiculo}, Código de cliente: {codigo_cliente}, Cantidad: {cantidad}, Total: {total}")

    def crear_venta(self, codigo_vehiculo, codigo_cliente, cantidad, total):
        query = "INSERT INTO ventas (codigo_vehiculo, codigo_cliente, cantidad, total) VALUES (%s, %s, %s, %s)"
        values = (codigo_vehiculo, codigo_cliente, cantidad, total)
        self.cursor.execute(query, values)
        self.conexion.commit()
        print("Venta registrada con éxito.")

    def anular_venta(self, codigo_vehiculo, codigo_cliente):
        query = "DELETE FROM ventas WHERE codigo_vehiculo = %s AND codigo_cliente = %s"
        values = (codigo_vehiculo, codigo_cliente)
        self.cursor.execute(query, values)
        self.conexion.commit()
        print("Venta anulada con éxito.")

    # Reportes Básicos
    def ventas_por_cliente(self, codigo_cliente):
        query = "SELECT codigo_vehiculo, cantidad, total FROM ventas WHERE codigo_cliente = %s"
        self.cursor.execute(query, (codigo_cliente,))
        print(f"Ventas del cliente (Código {codigo_cliente}):")
        for (codigo_vehiculo, cantidad, total) in self.cursor:
            print(f"Código de vehiculo: {codigo_vehiculo}, Cantidad: {cantidad}, Total: {total}")

    def ventas_por_vehiculo(self, codigo_vehiculo):
        query = "SELECT codigo_cliente, cantidad, total FROM ventas WHERE codigo_vehiculo = %s"
        self.cursor.execute(query, (codigo_vehiculo,))
        print(f"Ventas del vehiculo (Código {codigo_vehiculo}):")
        for (codigo_cliente, cantidad, total) in self.cursor:
            print(f"Código de cliente: {codigo_cliente}, Cantidad: {cantidad}, Total: {total}")

    def obtener_ventas_por_cliente(self, codigo_cliente):
        query = "SELECT codigo_vehiculo, codigo_cliente, cantidad, total FROM ventas WHERE codigo_cliente = %s"
        self.cursor.execute(query, (codigo_cliente,))
        return self.cursor.fetchall()

    def obtener_ventas_por_vehiculo(self, codigo_vehiculo):
        query = "SELECT codigo_vehiculo, codigo_cliente, cantidad, total FROM ventas WHERE codigo_vehiculo = %s"
        self.cursor.execute(query, (codigo_vehiculo,))
        return self.cursor.fetchall()

    def generar_reporte_ventas(self, ventas, nombre_reporte):
        doc = Document()
        doc.add_heading(f'Reporte de Ventas - ', 0)

        for venta in ventas:
            doc.add_paragraph(f'vehiculo: {venta[0]}')
            doc.add_paragraph(f'Cliente: {venta[1]}')
            doc.add_paragraph(f'Cantidad: {venta[2]}')
            doc.add_paragraph(f'Total: {venta[3]}')
            doc.add_paragraph('')

        doc.save(f'reporte_ventas.docx')
        print(f'Reporte generado y guardado como reporte_ventas.docx')


    def enviar_correo_con_adjunto(self, destinatario, asunto, mensaje, archivo_adjunto):
      
        servidor_smtp = "smtp.gmail.com" 
        puerto_smtp = 587  
        correo_emisor = "leoestuarlem@gmail.com"
        clave_emisor = ""  #Se elimino la contraseña por seguridad

        
        servidor = smtplib.SMTP(servidor_smtp, puerto_smtp)

      
        servidor.starttls()
        servidor.login(correo_emisor, clave_emisor)

       
        mensaje_correo = MIMEMultipart()
        mensaje_correo["From"] = correo_emisor
        mensaje_correo["To"] = destinatario
        mensaje_correo["Subject"] = asunto

       
        with open(archivo_adjunto, "rb") as adjunto:
            part = MIMEApplication(adjunto.read(), Name=os.path.basename(archivo_adjunto))
            part['Content-Disposition'] = f'attachment; filename="{os.path.basename(archivo_adjunto)}"'
            mensaje_correo.attach(part)

      
        mensaje_correo.attach(MIMEText(mensaje, "plain"))

       
        servidor.sendmail(correo_emisor, destinatario, mensaje_correo.as_string())

        servidor.quit()
    

def mostrar_menu():
    print("Menu:")
    print("1. Control de Inventario")
    print("2. Control de Clientes")
    print("3. Control de Ventas")
    print("4. Reportes Básicos")
    print("5. Salir")

def main():
    parser = argparse.ArgumentParser(description="Sistema de vehiculos")
    parser.add_argument("--ayuda", help="Mostrar ayuda", action="store_true")
    parser.add_argument("--inventario", help="Control de Inventario", action="store_true")
    parser.add_argument("--listar", help="Listar vehiculos", action="store_true")
    parser.add_argument("--crear", help="Crear vehiculo")
    parser.add_argument("--actualizar", help="Actualizar vehiculo")
    parser.add_argument("--existencia", help="Editar existencias de vehiculo")
    parser.add_argument("--eliminar", help="Eliminar vehiculo")

    args = parser.parse_args()
    sistema = SistemaVentas()

    if args.ayuda:
        parser.print_help()
    elif args.inventario:
        if args.listar:
            sistema.listar_vehiculo()
        elif args.crear:
            if len(args.crear) == 5:
                codigo, nombre, existencia, proveedor, precio = args.crear
                sistema.crear_vehiculo(int(codigo), nombre, int(existencia), proveedor, float(precio))
            else:
                print("Error: se requieren 5 argumentos para crear un vehiculo.")
        elif args.actualizar:
            codigo, nombre, existencia, proveedor, precio = args.actualizar.split()
            sistema.actualizar_vehiculo(codigo, nombre, int(existencia), proveedor, float(precio))
        elif args.existencia:
            codigo, cantidad = args.existencia.split()
            sistema.editar_existencias(codigo, int(cantidad))
        elif args.eliminar:
            codigo = args.eliminar
            sistema.eliminar_vehiculo(codigo)

            

     
 

    sistema = SistemaVentas()

    
    while True:
        mostrar_menu()
        opcion = input("Elija una opción: ")

        if opcion == "1":
            # Control de Inventario
            print("Control de vehiculos:")
            print("a. Listar vehiculos")
            print("b. Crear vehiculos")
            print("c. Actualizar vehiculo")
            print("d. Editar existencias de vehiculo")
            print("e. Eliminar vehiculo")
            sub_opcion = input("Elija una opción: ")

            if sub_opcion == "a":
                sistema.listar_vehiculo()
            elif sub_opcion == "b":
                codigo = input("Código: ")
                nombre = input("Nombre: ")
                existencia = int(input("Existencia: "))
                proveedor = input("Proveedor: ")
                precio = float(input("Precio: "))
                sistema.crear_vehiculo(codigo, nombre, existencia, proveedor, precio)
            elif sub_opcion == "c":
                codigo = input("Código del vehiculo a actualizar: ")
                nombre = input("Nuevo nombre: ")
                existencia = int(input("Nueva existencia: "))
                proveedor = input("Nuevo proveedor: ")
                precio = float(input("Nuevo precio: "))
                sistema.actualizar_vehiculo(codigo, nombre, existencia, proveedor, precio)
            elif sub_opcion == "d":
                codigo = input("Código del vehiculo a editar existencias: ")
                cantidad = int(input("Cantidad a agregar/reducir: "))
                sistema.editar_existencias(codigo, cantidad)
            elif sub_opcion == "e":
                codigo = input("Código del vehiculo a eliminar: ")
                sistema.eliminar_vehiculo(codigo)

        elif opcion == "2":
            # Control de Clientes
            print("Control de Clientes:")
            print("a. Listar clientes")
            print("b. Crear cliente")
            print("c. Editar cliente")
            print("d. Eliminar cliente")
            sub_opcion = input("Elija una opción: ")

            if sub_opcion == "a":
                sistema.listar_clientes()
            elif sub_opcion == "b":
                codigo = input("Código: ")
                nombre = input("Nombre: ")
                direccion = input("Dirección: ")
                sistema.crear_cliente(codigo, nombre, direccion)
            elif sub_opcion == "c":
                codigo = input("Código del cliente a editar: ")
                nombre = input("Nuevo nombre: ")
                direccion = input("Nueva dirección: ")
                sistema.editar_cliente(codigo, nombre, direccion)
            elif sub_opcion == "d":
                codigo = input("Código del cliente a eliminar: ")
                sistema.eliminar_cliente(codigo)


        elif opcion == "3":
            # Control de vehiculos
            print("Control de vehiculos:")
            print("a. Listar ventas")
            print("b. Crear venta")
            print("c. Anular venta")
            sub_opcion = input("Elija una opción: ")

            if sub_opcion == "a":
                sistema.listar_ventas()
            elif sub_opcion == "b":
                codigo_producto = input("Código de vehiculo: ")
                codigo_cliente = input("Código de cliente: ")
                cantidad = int(input("Cantidad de vehiculos: "))
                total = float(input("Total de venta: "))
                sistema.crear_venta(codigo_vehiculo, codigo_cliente, cantidad, total)
            elif sub_opcion == "c":
                codigo_vehiculo = input("Código del vehiculo de la venta a anular: ")
                codigo_cliente = input("Código del cliente de la venta a anular: ")
                sistema.anular_venta(codigo_vehiculo, codigo_cliente)
        elif opcion == "4":
            # Reportes Básicos
            print("Reportes Básicos:")
            print("a. Ventas por cliente")
            print("b. Ventas por vehiculo")
            sub_opcion = input("Elija una opción: ")

            if sub_opcion == "a":
                codigo_cliente = input("Código del cliente: ")
                sistema.ventas_por_cliente(codigo_cliente)
                ventas = sistema.obtener_ventas_por_cliente(codigo_cliente)
                sistema.generar_reporte_ventas(ventas, f'reporte_ventas')
                archivo_adjunto = f'reporte_ventas.docx'
                destinatario = '' #Se elimino el correo destinatario por seguridad 
                asunto = 'Reporte de Ventas por Cliente'
                mensaje = 'Adjunto se encuentra el reporte de ventas por cliente.'
                sistema.enviar_correo_con_adjunto(destinatario, asunto, mensaje, archivo_adjunto)
            elif sub_opcion == "b":
                codigo_vehiculo = input("Código del vehiculo: ")
                sistema.ventas_por_vehiculo(codigo_vehiculo)
                ventas = sistema.obtener_ventas_por_vehiculo(codigo_vehiculo)
                sistema.generar_reporte_ventas(ventas, f'reporte_ventas')
                archivo_adjunto = f'reporte_ventas.docx'
                destinatario = '' #Se elimino el correo destinatario por seguridad 
                asunto = 'Reporte de Ventas por vehiculo'
                mensaje = 'Adjunto se encuentra el reporte de ventas por vehiculo.'
                sistema.enviar_correo_con_adjunto(destinatario, asunto, mensaje, archivo_adjunto)
                

        elif opcion == "5":
            print("Saliendo del sistema de ventas.")
            sistema.cerrar_conexion()
            break
        else:
            print("Opción no válida. Intente de nuevo.")


if __name__ == "__main__":
    main()

